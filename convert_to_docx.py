#!/usr/bin/env python3
"""
Convert markdown documentation to DOCX format
"""
import sys
import re
import subprocess

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("Installing required package: python-docx")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

def create_document_styles(doc):
    """Create custom styles for the document"""
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 128)
    
    # Heading styles
    for i in range(1, 5):
        style_name = f'CustomHeading{i}'
        heading_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(20 - (i * 2))
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
    
    # Code style
    code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Table style
    table_style = doc.styles.add_style('TableText', WD_STYLE_TYPE.PARAGRAPH)
    table_style.font.name = 'Arial'
    table_style.font.size = Pt(10)
    
    return doc

def parse_markdown_table(lines):
    """Parse markdown table and return data"""
    if len(lines) < 3:
        return None
    
    # Extract headers
    headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    
    # Extract rows (skip separator line)
    rows = []
    for line in lines[2:]:
        if line.strip() and line.startswith('|'):
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(row)
        else:
            break
    
    return {'headers': headers, 'rows': rows, 'lines_consumed': len(rows) + 2}

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX format"""
    # Create document
    doc = Document()
    doc = create_document_styles(doc)
    
    # Set document properties
    doc.core_properties.title = "Crawl Infrastructure Documentation"
    doc.core_properties.author = "AI Assistant"
    doc.core_properties.subject = "Comprehensive documentation for the crawl infrastructure project"
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines for processing
    lines = content.split('\n')
    
    # Process lines
    i = 0
    in_code_block = False
    code_content = []
    in_list = False
    list_level = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph(code_text, style='CodeBlock')
                code_content = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            # Main title
            p = doc.add_paragraph(line[2:], style='CustomTitle')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_page_break()
            doc.add_paragraph(line[3:], style='CustomHeading1')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:], style='CustomHeading2')
        elif line.startswith('#### '):
            doc.add_paragraph(line[5:], style='CustomHeading3')
        
        # Tables
        elif line.strip().startswith('|') and i + 2 < len(lines) and lines[i + 1].strip().startswith('|'):
            # Parse table
            table_lines = [lines[j] for j in range(i, len(lines)) if lines[j].strip().startswith('|')]
            table_data = parse_markdown_table(table_lines)
            
            if table_data:
                # Create table
                table = doc.add_table(rows=1, cols=len(table_data['headers']))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                header_cells = table.rows[0].cells
                for j, header in enumerate(table_data['headers']):
                    header_cells[j].text = header
                    # Bold headers
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add rows
                for row_data in table_data['rows']:
                    row_cells = table.add_row().cells
                    for j, cell_data in enumerate(row_data[:len(table_data['headers'])]):
                        row_cells[j].text = cell_data
                
                i += table_data['lines_consumed']
                continue
        
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* ') or re.match(r'^\s*\d+\.\s', line):
            # Calculate indent level
            indent = len(line) - len(line.lstrip())
            current_level = indent // 2
            
            # Extract list item text
            if line.strip().startswith(('- ', '* ')):
                item_text = line.strip()[2:]
                list_style = 'List Bullet'
            else:
                # Numbered list
                item_text = re.sub(r'^\s*\d+\.\s*', '', line)
                list_style = 'List Number'
            
            # Handle inline formatting
            item_text = re.sub(r'\*\*(.+?)\*\*', r'\1', item_text)  # Bold
            item_text = re.sub(r'`(.+?)`', r'\1', item_text)  # Code
            
            p = doc.add_paragraph(item_text, style=list_style)
            if current_level > 0:
                p.paragraph_format.left_indent = Inches(0.5 * current_level)
        
        # Horizontal rule
        elif line.strip() == '---':
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Regular paragraphs
        elif line.strip():
            # Handle inline formatting
            text = line
            
            # Handle bold text
            parts = re.split(r'(\*\*.*?\*\*)', text)
            p = doc.add_paragraph()
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif '`' in part:
                    # Handle inline code
                    code_parts = re.split(r'(`.*?`)', part)
                    for code_part in code_parts:
                        if code_part.startswith('`') and code_part.endswith('`'):
                            run = p.add_run(code_part[1:-1])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 0, 139)
                        else:
                            p.add_run(code_part)
                else:
                    p.add_run(part)
        
        # Empty lines
        else:
            if not in_list:
                doc.add_paragraph()
        
        i += 1
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Crawl Infrastructure Documentation - Generated by AI Assistant"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(docx_file)
    print(f"DOCX file created: {docx_file}")

if __name__ == "__main__":
    md_file = "COMPREHENSIVE_DOCUMENTATION.md"
    docx_file = "CRAWL_INFRASTRUCTURE_DOCUMENTATION.docx"
    
    convert_markdown_to_docx(md_file, docx_file)